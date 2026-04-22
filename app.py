from flask import Flask, render_template, request, redirect, send_file, jsonify, url_for, session
from flask_wtf.csrf import CSRFProtect
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime, timedelta
import random
import json
import os
from pathlib import Path
from functools import wraps
import logging
from collections import defaultdict

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-change-in-production'
csrf = CSRFProtect(app)

# Admin password (HASHED) - change by modifying the hash
# To generate a new hash: from werkzeug.security import generate_password_hash; print(generate_password_hash('your_password'))
ADMIN_PASSWORD_HASH = generate_password_hash('admin123')

# Rate limiting: track failed login attempts
login_attempts = defaultdict(list)  # IP -> list of timestamps
MAX_LOGIN_ATTEMPTS = 5
LOCKOUT_DURATION = 300  # 5 minutes in seconds

# Setup logging
logging.basicConfig(filename='admin_login.log', level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')

# Create directory for storing generated  PDFs
EXAMS_DIR = Path("generated_exams")
EXAMS_DIR.mkdir(exist_ok=True)

# ========================
# DECORATORS & HELPERS
# ========================

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('admin_authenticated'):
            return redirect(url_for('admin_login'))
        # Verify IP hasn't changed (detect session hijacking)
        if session.get('admin_ip') != request.remote_addr:
            session.clear()
            logging.warning(f'Session hijacking attempt detected: original IP {session.get("admin_ip")}, current IP {request.remote_addr}')
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated_function

def is_rate_limited(ip):
    """Check if IP is rate limited"""
    now = datetime.now().timestamp()
    # Remove old attempts outside the rate limit window
    login_attempts[ip] = [ts for ts in login_attempts[ip] if now - ts < LOCKOUT_DURATION]
    return len(login_attempts[ip]) >= MAX_LOGIN_ATTEMPTS

def record_login_attempt(ip, success=False):
    """Record a login attempt"""
    now = datetime.now().timestamp()
    if not success:
        login_attempts[ip].append(now)
    else:
        # Clear attempts on successful login
        login_attempts[ip] = []
    
    status = 'SUCCESS' if success else 'FAILED'
    logging.info(f'Admin login {status}: IP={ip}')

def get_db():
    conn = sqlite3.connect("testbank.db")
    conn.row_factory = sqlite3.Row
    return conn

def shuffle_and_relabel_choices(choices_list, correct_answer_letter):
    """Shuffle choices and reassign labels A, B, C, D, updating correct answer"""
    # Convert sqlite3.Row objects to dictionaries for modification
    choices = [dict(choice) for choice in choices_list]
    
    # Track which choice object is the correct one
    correct_choice_obj = None
    for choice in choices:
        if choice['choice_label'] == correct_answer_letter:
            correct_choice_obj = choice
            break
    
    # Shuffle the choices
    random.shuffle(choices)
    
    # Reassign labels A, B, C, D and find new correct answer
    new_correct_answer = None
    for idx, choice in enumerate(choices):
        new_label = chr(65 + idx)  # A=65, B=66, C=67, D=68
        choice['choice_label'] = new_label
        if choice is correct_choice_obj:
            new_correct_answer = new_label
    
    return choices, new_correct_answer

def sanitize_filename(filename):
    """Remove or replace invalid filename characters"""
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    return filename.strip()

def save_exam_to_history(conn, chapter_id, filename, file_path, num_questions, include_problems):
    """Record generated exam in exam_history table"""
    cursor = conn.cursor()
    # Use local system time instead of UTC
    created_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    cursor.execute("""
        INSERT INTO exam_history (chapter_id, filename, file_path, created_date, num_questions, include_problems)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (chapter_id, filename, str(file_path), created_date, num_questions, 1 if include_problems else 0))
    conn.commit()
    return cursor.lastrowid

def generate_docx_exam(chapter, questions_with_choices, include_problems):
    """Generate exam as DOCX document"""
    doc = Document()
    
    # Title and chapter info
    title = doc.add_heading(f"Chapter {chapter['chapter_number']}: {chapter['chapter_title']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    exam_heading = doc.add_heading('Examination', level=2)
    exam_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Student info section
    student_info = doc.add_paragraph()
    student_info.add_run('Student Name: ').bold = True
    student_info.add_run('_' * 50)
    student_info.add_run('     Date: ').bold = True
    student_info.add_run('_' * 20)
    
    doc.add_paragraph()  # Blank line
    
    # Questions
    for idx, (question, choices) in enumerate(questions_with_choices, 1):
        # Question number if included
        if include_problems and question.get('question_number'):
            qnum_para = doc.add_paragraph()
            qnum_run = qnum_para.add_run(f'Question {question["question_number"]}')
            qnum_run.italic = True
            qnum_run.font.color.rgb = RGBColor(13, 115, 119)  # Teal color
        
        # Question text
        question_para = doc.add_paragraph(style='List Number')
        question_para.paragraph_format.left_indent = Inches(0)
        question_para.clear()
        question_run = question_para.add_run(f'{idx}. {question["question_text"]}')
        question_run.bold = True
        
        # Choice options
        for choice in choices:
            choice_para = doc.add_paragraph(
                f'{choice["choice_label"]}. {choice["choice_text"]}',
                style='List Bullet'
            )
            choice_para.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_paragraph()  # Blank line between questions
    
    return doc

@app.route("/")
def index():
    """Dashboard - choose REVIEW or EXAM mode"""
    return render_template("mode_selector.html")

@app.route("/admin/login", methods=['GET', 'POST'])
def admin_login():
    """Admin password login"""
    ip = request.remote_addr
    
    if request.method == 'POST':
        # Check rate limiting
        if is_rate_limited(ip):
            record_login_attempt(ip, success=False)
            return render_template('admin_login.html', error='🔒 Too many failed attempts. Please try again in 5 minutes.')
        
        password = request.form.get('password', '')
        
        if check_password_hash(ADMIN_PASSWORD_HASH, password):
            session['admin_authenticated'] = True
            session['admin_ip'] = ip  # Store IP for session hijacking detection
            record_login_attempt(ip, success=True)
            return redirect(url_for('manage_subjects'))
        else:
            record_login_attempt(ip, success=False)
            attempts_left = MAX_LOGIN_ATTEMPTS - len([ts for ts in login_attempts[ip] if datetime.now().timestamp() - ts < LOCKOUT_DURATION])
            return render_template('admin_login.html', error=f'❌ Incorrect password. {attempts_left} attempts remaining.')
    
    return render_template('admin_login.html')

@app.route("/admin/logout")
def admin_logout():
    """Admin logout"""
    ip = request.remote_addr
    logging.info(f'Admin logout: IP={ip}')
    session.clear()
    return redirect('/')

@app.route("/admin/change-password", methods=['GET', 'POST'])
def change_password():
    """Change admin password"""
    global ADMIN_PASSWORD_HASH   # ✅ MOVE IT HERE

    if request.method == 'POST':
        current_password = request.form.get('current_password', '')
        new_password = request.form.get('new_password', '')
        confirm_password = request.form.get('confirm_password', '')
        
        # Validate current password
        if not check_password_hash(ADMIN_PASSWORD_HASH, current_password):
            return render_template('admin_change_password.html', error='❌ Current password is incorrect.')
        
        # Validate new password length
        if len(new_password) < 8:
            return render_template('admin_change_password.html', error='❌ New password must be at least 8 characters long.')
        
        # Check passwords match
        if new_password != confirm_password:
            return render_template('admin_change_password.html', error='❌ New passwords do not match.')
        
        # Update password hash
        ADMIN_PASSWORD_HASH = generate_password_hash(new_password)
        
        # Log the password change
        logging.info(f'Admin password changed: IP={request.remote_addr}')
        
        return render_template('admin_change_password.html', success=True)
    
    return render_template('admin_change_password.html')

# ========================
# SUBJECTS & TOPICS MANAGEMENT
# ========================

@app.route("/api/subjects")
def get_subjects():
    """Get all subjects with their topics and chapters"""
    conn = get_db()
    cursor = conn.cursor()
    
    cursor.execute("SELECT id, name, description FROM subjects ORDER BY name")
    subjects = cursor.fetchall()
    
    result = []
    for subject in subjects:
        subject_dict = dict(subject)
        
        # Get topics for this subject
        cursor.execute("""
            SELECT id, name, description FROM topics 
            WHERE subject_id = ? ORDER BY name
        """, (subject['id'],))
        topics = cursor.fetchall()
        
        subject_dict['topics'] = []
        for topic in topics:
            topic_dict = dict(topic)
            
            # Get chapters for this topic
            cursor.execute("""
                SELECT id, chapter_number, chapter_title FROM chapters 
                WHERE topic_id = ? ORDER BY chapter_number
            """, (topic['id'],))
            chapters = cursor.fetchall()
            topic_dict['chapters'] = [dict(ch) for ch in chapters]
            subject_dict['topics'].append(topic_dict)
        
        result.append(subject_dict)
    
    conn.close()
    return jsonify(result)

@app.route("/admin/subjects", methods=['GET', 'POST'])
@admin_required
def manage_subjects():
    """Admin: manage subjects and topics"""
    conn = get_db()
    cursor = conn.cursor()
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add_subject':
            name = request.form.get('subject_name')
            description = request.form.get('subject_description', '')
            cursor.execute(
                "INSERT INTO subjects (name, description) VALUES (?, ?)",
                (name, description)
            )
            conn.commit()
        
        elif action == 'add_topic':
            subject_id = request.form.get('subject_id')
            name = request.form.get('topic_name')
            description = request.form.get('topic_description', '')
            cursor.execute(
                "INSERT INTO topics (subject_id, name, description) VALUES (?, ?, ?)",
                (subject_id, name, description)
            )
            conn.commit()
        
        elif action == 'edit_subject':
            subject_id = request.form.get('subject_id')
            name = request.form.get('subject_name')
            description = request.form.get('subject_description', '')
            cursor.execute(
                "UPDATE subjects SET name = ?, description = ? WHERE id = ?",
                (name, description, subject_id)
            )
            conn.commit()
        
        elif action == 'edit_topic':
            topic_id = request.form.get('topic_id')
            name = request.form.get('topic_name')
            description = request.form.get('topic_description', '')
            cursor.execute(
                "UPDATE topics SET name = ?, description = ? WHERE id = ?",
                (name, description, topic_id)
            )
            conn.commit()
        
        elif action == 'delete_subject':
            subject_id = request.form.get('subject_id')
            cursor.execute("DELETE FROM subjects WHERE id = ?", (subject_id,))
            conn.commit()
        
        elif action == 'delete_topic':
            topic_id = request.form.get('topic_id')
            cursor.execute("DELETE FROM topics WHERE id = ?", (topic_id,))
            conn.commit()
    
    # Fetch all subjects with their topics
    cursor.execute("SELECT id, name, description FROM subjects ORDER BY name")
    subjects_rows = cursor.fetchall()
    
    subjects_data = []
    for subject in subjects_rows:
        subject_dict = dict(subject)
        
        # Get topics for this subject
        cursor.execute(
            "SELECT id, name, description FROM topics WHERE subject_id = ? ORDER BY name",
            (subject['id'],)
        )
        topics = cursor.fetchall()
        subject_dict['topics'] = [dict(t) for t in topics]
        subjects_data.append(subject_dict)
    
    conn.close()
    
    return render_template("admin_subjects.html", subjects=subjects_data)

# ========================
# CHAPTERS MANAGEMENT
# ========================

@app.route("/admin/chapters/<int:topic_id>", methods=['GET', 'POST'])
@admin_required
def manage_chapters(topic_id):
    """Manage chapters for a topic"""
    conn = get_db()
    cursor = conn.cursor()
    
    # Get topic info
    cursor.execute("SELECT id, name, subject_id FROM topics WHERE id = ?", (topic_id,))
    topic = cursor.fetchone()
    
    if not topic:
        conn.close()
        return redirect(url_for('manage_subjects'))
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add_chapter':
            chapter_number = request.form.get('chapter_number')
            chapter_title = request.form.get('chapter_title')
            cursor.execute(
                "INSERT INTO chapters (topic_id, chapter_number, chapter_title) VALUES (?, ?, ?)",
                (topic_id, chapter_number, chapter_title)
            )
            conn.commit()
        
        elif action == 'edit_chapter':
            chapter_id = request.form.get('chapter_id')
            chapter_number = request.form.get('chapter_number')
            chapter_title = request.form.get('chapter_title')
            cursor.execute(
                "UPDATE chapters SET chapter_number = ?, chapter_title = ? WHERE id = ?",
                (chapter_number, chapter_title, chapter_id)
            )
            conn.commit()
        
        elif action == 'delete_chapter':
            chapter_id = request.form.get('chapter_id')
            cursor.execute("DELETE FROM chapters WHERE id = ?", (chapter_id,))
            conn.commit()
    
    # Fetch all chapters for this topic
    cursor.execute(
        "SELECT id, chapter_number, chapter_title FROM chapters WHERE topic_id = ? ORDER BY chapter_number",
        (topic_id,)
    )
    chapters = cursor.fetchall()
    
    conn.close()
    
    return render_template("admin_chapters.html", topic=dict(topic), chapters=[dict(c) for c in chapters])

# ========================
# QUESTIONS MANAGEMENT
# ========================

@app.route("/admin/question_groups/<int:chapter_id>", methods=['GET', 'POST'])
@admin_required
def manage_question_groups(chapter_id):
    """Manage question groups (question numbers) for a chapter"""
    conn = get_db()
    cursor = conn.cursor()
    
    # Get chapter info
    cursor.execute("""
        SELECT c.id, c.chapter_number, c.chapter_title, t.id as topic_id, t.name as topic_name, s.id as subject_id, s.name as subject_name
        FROM chapters c
        JOIN topics t ON c.topic_id = t.id
        JOIN subjects s ON t.subject_id = s.id
        WHERE c.id = ?
    """, (chapter_id,))
    chapter = cursor.fetchone()
    
    if not chapter:
        conn.close()
        return redirect(url_for('manage_subjects'))
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add_question_group':
            question_number = request.form.get('question_number')
            section = request.form.get('section', '')
            standard = request.form.get('standard', '')
            
            cursor.execute(
                "INSERT INTO question_groups (chapter_id, question_number, section, standard) VALUES (?, ?, ?, ?)",
                (chapter_id, question_number, section, standard)
            )
            conn.commit()
        
        elif action == 'edit_question_group':
            qg_id = request.form.get('question_group_id')
            question_number = request.form.get('question_number')
            section = request.form.get('section', '')
            standard = request.form.get('standard', '')
            
            cursor.execute(
                "UPDATE question_groups SET question_number = ?, section = ?, standard = ? WHERE id = ?",
                (question_number, section, standard, qg_id)
            )
            conn.commit()
        
        elif action == 'delete_question_group':
            qg_id = request.form.get('question_group_id')
            cursor.execute("DELETE FROM question_groups WHERE id = ?", (qg_id,))
            conn.commit()
    
    # Fetch all question groups with their question count for this chapter
    cursor.execute("""
        SELECT qg.id, qg.question_number, qg.section, qg.standard, COUNT(q.id) as question_count
        FROM question_groups qg
        LEFT JOIN questions q ON qg.id = q.question_group_id
        WHERE qg.chapter_id = ?
        GROUP BY qg.id
        ORDER BY qg.question_number
    """, (chapter_id,))
    question_groups = cursor.fetchall()
    
    conn.close()
    
    return render_template("admin_question_groups.html", chapter=dict(chapter), question_groups=[dict(qg) for qg in question_groups])

@app.route("/admin/questions/<int:question_group_id>", methods=['GET', 'POST'])
@admin_required
def manage_questions(question_group_id):
    """Manage questions for a question group"""
    conn = get_db()
    cursor = conn.cursor()
    
    # Get question group and chapter info
    cursor.execute("""
        SELECT qg.id, qg.question_number, qg.section, qg.standard, 
               c.id as chapter_id, c.chapter_number, c.chapter_title, t.id as topic_id, t.name as topic_name
        FROM question_groups qg
        JOIN chapters c ON qg.chapter_id = c.id
        JOIN topics t ON c.topic_id = t.id
        WHERE qg.id = ?
    """, (question_group_id,))
    question_group = cursor.fetchone()
    
    if not question_group:
        conn.close()
        return redirect(url_for('manage_subjects'))
    
    if request.method == 'POST':
        action = request.form.get('action')
        
        if action == 'add_question':
            question_text = request.form.get('question_text')
            question_type = request.form.get('question_type', 'multiple_choice')
            correct_choice = request.form.get('correct_choice')
            explanation = request.form.get('explanation', '')
            
            cursor.execute(
                "INSERT INTO questions (question_group_id, question_text, correct_choice, explanation, question_type) VALUES (?, ?, ?, ?, ?)",
                (question_group_id, question_text, correct_choice, explanation, question_type)
            )
            question_id = cursor.lastrowid
            conn.commit()
            
            # Add choices for multiple choice questions
            if question_type == 'multiple_choice':
                for i in range(4):
                    choice_label = chr(65 + i)  # A, B, C, D
                    choice_text = request.form.get(f'choice_{choice_label}', '')
                    if choice_text:
                        cursor.execute(
                            "INSERT INTO choices (question_id, choice_label, choice_text) VALUES (?, ?, ?)",
                            (question_id, choice_label, choice_text)
                        )
                conn.commit()
        
        elif action == 'edit_question':
            question_id = request.form.get('question_id')
            question_text = request.form.get('question_text')
            question_type = request.form.get('question_type', 'multiple_choice')
            correct_choice = request.form.get('correct_choice')
            explanation = request.form.get('explanation', '')
            
            cursor.execute(
                "UPDATE questions SET question_text = ?, correct_choice = ?, explanation = ?, question_type = ? WHERE id = ?",
                (question_text, correct_choice, explanation, question_type, question_id)
            )
            conn.commit()
            
            # Update choices for multiple choice questions
            if question_type == 'multiple_choice':
                # Delete old choices
                cursor.execute("DELETE FROM choices WHERE question_id = ?", (question_id,))
                conn.commit()
                
                # Add new choices
                for i in range(4):
                    choice_label = chr(65 + i)  # A, B, C, D
                    choice_text = request.form.get(f'choice_{choice_label}', '')
                    if choice_text:
                        cursor.execute(
                            "INSERT INTO choices (question_id, choice_label, choice_text) VALUES (?, ?, ?)",
                            (question_id, choice_label, choice_text)
                        )
                conn.commit()
        
            conn.commit()
    
    # Fetch all questions for this question group
    cursor.execute("""
        SELECT q.id, q.question_text, q.question_type, q.correct_choice, q.explanation
        FROM questions q
        WHERE q.question_group_id = ?
        ORDER BY q.id
    """, (question_group_id,))
    questions = cursor.fetchall()
    
    conn.close()
    
    return render_template("admin_questions.html", question_group=dict(question_group), questions=[dict(q) for q in questions])

@app.route("/api/question/<int:question_id>")
def get_question_details(question_id):
    """Get question details including choices"""
    conn = get_db()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT id, question_text, question_type, correct_choice, explanation
        FROM questions
        WHERE id = ?
    """, (question_id,))
    question = cursor.fetchone()
    
    if not question:
        conn.close()
        return jsonify({'error': 'Question not found'}), 404
    
    q_dict = dict(question)
    
    # Get choices if multiple choice
    if q_dict['question_type'] == 'multiple_choice':
        cursor.execute(
            "SELECT choice_label, choice_text FROM choices WHERE question_id = ? ORDER BY choice_label",
            (question_id,)
        )
        choices = cursor.fetchall()
        q_dict['choices'] = [dict(c) for c in choices]
    
    conn.close()
    return jsonify(q_dict)

# ========================
# REVIEW MODE
# ========================

@app.route("/review")
def review_mode():
    """REVIEW mode - select subject/topic/chapter"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM subjects ORDER BY name")
    subjects = cursor.fetchall()
    conn.close()
    
    return render_template("review_selector.html", subjects=subjects)

@app.route("/api/topics/<int:subject_id>")
def get_topics(subject_id):
    """Get topics for a subject"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, name FROM topics WHERE subject_id = ? ORDER BY name",
        (subject_id,)
    )
    topics = cursor.fetchall()
    conn.close()
    return jsonify([dict(t) for t in topics])

@app.route("/api/chapters/<int:topic_id>")
def get_chapters(topic_id):
    """Get chapters for a topic"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, chapter_number, chapter_title FROM chapters WHERE topic_id = ? ORDER BY chapter_number",
        (topic_id,)
    )
    chapters = cursor.fetchall()
    conn.close()
    return jsonify([dict(c) for c in chapters])

@app.route("/review/start", methods=['POST'])
def start_review():
    """Start a review session"""
    chapter_id = request.form.get('chapter_id')
    num_questions = int(request.form.get('num_questions', 10))
    
    conn = get_db()
    cursor = conn.cursor()
    
    # Get chapter info
    cursor.execute(
        "SELECT id, chapter_number, chapter_title FROM chapters WHERE id = ?",
        (chapter_id,)
    )
    chapter = cursor.fetchone()
    
    # Get random questions with choices (all types)
    cursor.execute("""
        SELECT q.id, q.question_text, q.correct_choice, q.explanation, q.question_type,
               qg.question_number
        FROM questions q
        JOIN question_groups qg ON q.question_group_id = qg.id
        WHERE qg.chapter_id = ?
        ORDER BY RANDOM()
        LIMIT ?
    """, (chapter_id, num_questions))
    
    questions = cursor.fetchall()
    
    # Get choices for each question and shuffle
    questions_data = []
    for question in questions:
        q_dict = dict(question)
        
        cursor.execute(
            "SELECT choice_label, choice_text FROM choices WHERE question_id = ? ORDER BY choice_label",
            (question['id'],)
        )
        choices = cursor.fetchall()
        choices_list = [dict(c) for c in choices]
        
        # Shuffle choices for multiple choice questions only
        if q_dict['question_type'] == 'multiple_choice' and choices_list:
            choices_list, q_dict['correct_choice'] = shuffle_and_relabel_choices(
                choices_list, q_dict['correct_choice']
            )
        
        q_dict['choices'] = choices_list
        questions_data.append(q_dict)
    
    conn.close()
    
    # Store session data
    session_data = {
        'chapter': dict(chapter) if chapter else {},
        'questions': questions_data
    }
    
    return render_template("review_exam.html", session=session_data)

# ========================
# EXAM MODE
# ========================

@app.route("/exam")
def exam_mode():
    """EXAM mode - select subject/topic/chapter"""
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM subjects ORDER BY name")
    subjects = cursor.fetchall()
    conn.close()
    
    return render_template("exam_selector.html", subjects=subjects)

@app.route("/exam/start", methods=['POST'])
def start_exam():
    """Start an exam session"""
    chapter_id = request.form.get('chapter_id')
    num_questions = int(request.form.get('num_questions', 20))
    total_time_limit = int(request.form.get('total_time_limit', 3600))  # in seconds
    
    conn = get_db()
    cursor = conn.cursor()
    
    # Get chapter info
    cursor.execute(
        "SELECT id, chapter_number, chapter_title FROM chapters WHERE id = ?",
        (chapter_id,)
    )
    chapter = cursor.fetchone()
    
    # Get random questions (all types - will only grade multiple choice)
    cursor.execute("""
        SELECT q.id, q.question_text, q.correct_choice, q.explanation, q.question_type,
               qg.question_number
        FROM questions q
        JOIN question_groups qg ON q.question_group_id = qg.id
        WHERE qg.chapter_id = ?
        ORDER BY RANDOM()
        LIMIT ?
    """, (chapter_id, num_questions))
    
    questions = cursor.fetchall()
    
    # Get choices for each question and shuffle
    questions_data = []
    for question in questions:
        q_dict = dict(question)
        
        cursor.execute(
            "SELECT choice_label, choice_text FROM choices WHERE question_id = ? ORDER BY choice_label",
            (question['id'],)
        )
        choices = cursor.fetchall()
        choices_list = [dict(c) for c in choices]
        
        # Shuffle choices for multiple choice questions only
        if q_dict['question_type'] == 'multiple_choice' and choices_list:
            choices_list, q_dict['correct_choice'] = shuffle_and_relabel_choices(
                choices_list, q_dict['correct_choice']
            )
        
        q_dict['choices'] = choices_list
        questions_data.append(q_dict)
    
    conn.close()
    
    # Store session data
    session_data = {
        'chapter': dict(chapter) if chapter else {},
        'questions': questions_data,
        'total_time_limit': total_time_limit
    }
    
    return render_template("exam_interface.html", session=session_data)

@app.route("/exam/submit", methods=['POST'])
def submit_exam():
    """Submit exam answers and calculate score"""
    data = request.get_json()
    answers = data.get('answers', {})  # {question_id: selected_choice or text}
    session_data = data.get('session', {})
    
    questions = session_data.get('questions', [])
    
    score = 0
    total_mc = 0  # Total multiple choice questions
    results = []
    
    for question in questions:
        q_id = str(question['id'])
        
        if question['question_type'] == 'multiple_choice':
            total_mc += 1
            selected = answers.get(q_id)
            correct = question['correct_choice']
            
            is_correct = selected == correct
            if is_correct:
                score += 1
            
            results.append({
                'id': question['id'],
                'type': 'multiple_choice',
                'question_text': question['question_text'],
                'selected': selected,
                'correct': correct,
                'is_correct': is_correct,
                'explanation': question.get('explanation', '')
            })
        else:
            # Descriptive question
            user_answer = answers.get(q_id, '')
            results.append({
                'id': question['id'],
                'type': 'descriptive',
                'question_text': question['question_text'],
                'user_answer': user_answer,
                'model_answer': question.get('correct_choice', 'N/A'),
                'explanation': question.get('explanation', ''),
                'is_correct': None  # Not auto-graded
            })
    
    percentage = (score / total_mc * 100) if total_mc > 0 else 0
    
    return jsonify({
        'score': score,
        'total': total_mc,
        'percentage': percentage,
        'results': results
    })

if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5000, use_reloader=False)
